
import os
import asyncio
import base64
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches, Pt

# Paths to screenshots
SCREENSHOT_PATHS = {
    "homepage": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\homepage_full_1775530147303.png",
    "policies": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\policies_page_1775530167123.png",
    "calendar": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\calendar_page_1775530182383.png",
    "projects": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\projects_page_1775530197769.png",
    "archives": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\project_archives_page_1775530427394.png",
    "knowledge": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\knowledge_products_page_1775530221301.png",
    "brochures": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\brochures_page_1775530246284.png",
    "livelihood": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\livelihood_program_page_1775530289469.png",
    "vision": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\vision_mission_page_1775530319392.png",
    "org": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\org_structure_page_1775530337669.png",
    "committee": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\gad_committee_page_1775530410091.png"
}

def get_base64(path):
    if os.path.exists(path):
        with open(path, "rb") as f:
            return f"data:image/png;base64,{base64.b64encode(f.read()).decode('utf-8')}"
    return ""

SCREENSHOTS = {k: get_base64(v) for k, v in SCREENSHOT_PATHS.items()}

HTML_CONTENT = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>GAD Corner Citizen's Guide</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&family=Outfit:wght@400;700&display=swap" rel="stylesheet">
    <style>
        :root {{
            --primary: #10b981;
            --secondary: #3b82f6;
            --dark: #1f2937;
            --text: #4b5563;
            --light: #f9fafb;
            --border: #e5e7eb;
        }}
        body {{ font-family: 'Inter', sans-serif; color: var(--text); margin: 0; padding: 0; }}
        .page {{ padding: 1in; page-break-after: always; position: relative; min-height: 10in; box-sizing: border-box; }}
        .page::before {{ content: ''; position: absolute; top: 0; left: 0; width: 6pt; height: 100%; background: var(--primary); }}
        h1, h2, h3 {{ font-family: 'Outfit', sans-serif; color: var(--dark); }}
        h1 {{ font-size: 42pt; margin-top: 2in; text-align: center; }}
        h2 {{ font-size: 22pt; border-bottom: 2pt solid var(--border); padding-bottom: 10pt; margin-top: 40pt; color: var(--primary); }}
        h3 {{ font-size: 16pt; margin-top: 20pt; }}
        .figure {{ margin: 25pt 0; text-align: center; page-break-inside: avoid; }}
        .figure img {{ width: 90%; border-radius: 10pt; box-shadow: 0 4pt 15pt rgba(0,0,0,0.05); border: 1pt solid var(--border); }}
        .figure-caption {{ font-size: 9pt; color: #9ca3af; margin-top: 8pt; }}
        p {{ line-height: 1.6; margin-bottom: 12pt; }}
        .info-box {{ background: #ecfdf5; border-left: 4pt solid var(--primary); padding: 15pt; border-radius: 8pt; margin: 20pt 0; }}
    </style>
</head>
<body>
    <div class="page" style="text-align: center; display: flex; flex-direction: column; justify-content: center;">
        <div style="font-weight: 700; color: var(--primary); letter-spacing: 0.1em; text-transform: uppercase;">Public Manual</div>
        <h1>GAD Corner</h1>
        <p style="font-size: 1.5rem;">Citizen's User Guide</p>
        <p>A comprehensive guide to exploring the Gender and Development (GAD) initiatives of Montalban, Rizal.</p>
    </div>

    <div class="page">
        <h2>1. Overview</h2>
        <p>Welcome to the <strong>GAD Corner</strong>, your portal for information, resources, and updates regarding Gender and Development in our municipality. This guide will help you navigate the platform to access policies, reports, and community projects.</p>
        
        <div class="figure">
            <img src="{SCREENSHOTS['homepage']}" alt="Homepage">
            <div class="figure-caption">Figure 1: The GAD Corner Homepage</div>
        </div>
        <p>The Homepage features the latest news and announcements. Use the navigation bar at the top to visit different sections of the site.</p>
    </div>

    <div class="page">
        <h2>2. Accessing GAD Policies</h2>
        <p>Stay informed about your rights and local mandates by exploring the Policies section. We archive Circulars, Memos, and Local Policies for public transparency.</p>
        
        <div class="figure">
            <img src="{SCREENSHOTS['policies']}" alt="Policies">
            <div class="figure-caption">Figure 2: Digital Policy Repository</div>
        </div>
        <p>You can browse documents by year and category. Click on any document to view its details or download the official PDF file.</p>
    </div>

    <div class="page">
        <h2>3. Community Calendar</h2>
        <p>Never miss a GAD event! Our interactive calendar displays all upcoming workshops, seminars, and community activities.</p>
        
        <div class="figure">
            <img src="{SCREENSHOTS['calendar']}" alt="Calendar">
            <div class="figure-caption">Figure 3: GAD Community Calendar</div>
        </div>
        <p>Click on an event to see more information, including time, location, and the type of activity (Community, Governance, etc.).</p>
    </div>

    <div class="page">
        <h2>4. Projects & Impact</h2>
        <p>Transparency is our priority. In the Projects Archive, you can track the status of GAD initiatives in our municipality.</p>
        
        <div class="figure">
            <img src="{SCREENSHOTS['projects']}" alt="Projects">
            <div class="figure-caption">Figure 4: Active Projects and Gallery</div>
        </div>
        
        <div class="figure">
            <img src="{SCREENSHOTS['archives']}" alt="Archives">
            <div class="figure-caption">Figure 5: Historical Project Archives</div>
        </div>
        <p>See high-resolution photos and detailed descriptions of both ongoing and completed projects.</p>
    </div>

    <div class="page">
        <h2>5. Resource Center</h2>
        <p>Our Resource Center provides free learning materials, toolkits, and brochures for all citizens.</p>
        
        <div class="figure">
            <img src="{SCREENSHOTS['knowledge']}" alt="Knowledge">
            <div class="figure-caption">Figure 6: Knowledge Products & Manuals</div>
        </div>
        
        <div class="figure">
            <img src="{SCREENSHOTS['brochures']}" alt="Brochures">
            <div class="figure-caption">Figure 7: Downloadable Brochures</div>
        </div>
    </div>

    <div class="page">
        <h2>6. Livelihood Program</h2>
        <p>We link direct feeds from our social media and external platforms to keep you updated on livelihood training and opportunities.</p>
        
        <div class="figure">
            <img src="{SCREENSHOTS['livelihood']}" alt="Livelihood">
            <div class="figure-caption">Figure 8: Livelihood Announcements and Feeds</div>
        </div>
    </div>

    <div class="page">
        <h2>7. About the GFPS</h2>
        <p>Learn about the GAD Focal Point System (GFPS), our organizational structure, and the dedicated members of our committee.</p>
        
        <div class="figure">
            <img src="{SCREENSHOTS['vision']}" alt="Vision">
            <div class="figure-caption">Figure 9: Vision & Mission of GAD Corner</div>
        </div>
        
        <div class="figure">
            <img src="{SCREENSHOTS['org']}" alt="Org Structure">
            <div class="figure-caption">Figure 10: Organizational Hierarchy</div>
        </div>

        <div class="figure">
            <img src="{SCREENSHOTS['committee']}" alt="Committee">
            <div class="figure-caption">Figure 11: Meet the GAD Committee</div>
        </div>
    </div>

    <div class="page" style="text-align: center; justify-content: center; display: flex; flex-direction: column;">
        <h2>Thank You for Visiting</h2>
        <p>Gender equality and community empowerment start with informed citizens.</p>
        <div style="margin-top: 40px; font-size: 0.8rem; color: #9ca3af;">
            &copy; 2026 GAD Corner | Municipality of Montalban
        </div>
    </div>
</body>
</html>
"""

async def generate_pdf():
    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page()
        await page.set_content(HTML_CONTENT)
        await page.wait_for_load_state("networkidle")
        await page.wait_for_timeout(3000)
        
        pdf_path = os.path.join(os.getcwd(), "GAD_Citizen_Guide.pdf")
        await page.pdf(
            path=pdf_path, 
            format="A4", 
            print_background=True,
            margin={"top": "40px", "right": "20px", "bottom": "40px", "left": "20px"}
        )
        await browser.close()
        print(f"PDF generated: {pdf_path}")

def generate_docx():
    from docx import Document
    doc = Document()
    doc.add_heading('GAD Corner - Citizen Guide', 0)
    
    sections = [
        ("1. Overview", SCREENSHOT_PATHS['homepage'], "The GAD Corner Homepage"),
        ("2. Policies", SCREENSHOT_PATHS['policies'], "Digital Policy Repository"),
        ("3. Calendar", SCREENSHOT_PATHS['calendar'], "Community Calendar"),
        ("4. Projects", SCREENSHOT_PATHS['projects'], "Active Projects"),
        ("5. Knowledge Products", SCREENSHOT_PATHS['knowledge'], "Learning materials"),
        ("6. Livelihood", SCREENSHOT_PATHS['livelihood'], "Social feeds"),
        ("7. GFPS Structure", SCREENSHOT_PATHS['org'], "Organizational structure")
    ]
    
    for title, img, cap in sections:
        doc.add_heading(title, level=1)
        if os.path.exists(img):
            doc.add_picture(img, width=Inches(6))
            p = doc.add_paragraph(cap)
            p.alignment = 1
        doc.add_page_break()
        
    docx_path = os.path.join(os.getcwd(), "GAD_Citizen_Guide.docx")
    doc.save(docx_path)
    print(f"DOCX generated: {docx_path}")

if __name__ == "__main__":
    asyncio.run(generate_pdf())
    generate_docx()
