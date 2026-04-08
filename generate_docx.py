
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Screenshots paths (same as before)
SCREENSHOT_PATHS = {
    "dashboard": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_dashboard_1775529097427.png",
    "features": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_features_1775529107736.png",
    "policies": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_policies_1775529119559.png",
    "events": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_events_1775529130074.png",
    "projects": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_projects_1775529141742.png",
    "knowledge": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_knowledge_1775529152781.png",
    "brochures": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_brochures_1775529164771.png",
    "livelihood": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_livelihood_feeds_1775529176466.png",
    "org": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_org_structure_1775529188145.png",
    "committee": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_committee_1775529199795.png",
    "tracking": r"C:\Users\Administrator\.gemini\antigravity\brain\f13b0cfb-6def-4e14-815d-a5e87357b6c8\admin_tracking_matrix_1775529211409.png"
}

def create_manual():
    doc = Document()
    
    # --- Title Page ---
    title = doc.add_heading('GAD Corner', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Administrative User Manual\nMunicipality of Montalban (Rodriguez), Rizal')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph("Welcome to the GAD Corner Admin Panel. This user manual will guide you through managing the Gender and Development platform, from content updates to administrative oversight.")
    
    # --- 2. Authentication ---
    doc.add_heading('2. Authentication', level=1)
    doc.add_paragraph("Access the admin panel at /admin/login. Only authorized officers should possess credentials.")
    if os.path.exists(SCREENSHOT_PATHS["dashboard"]):
        doc.add_picture(SCREENSHOT_PATHS["dashboard"], width=Inches(6))
        cap = doc.add_paragraph("Figure 1: Admin Portal Login Dashboard Overview")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 3. Dashboard Overview ---
    doc.add_heading('3. Dashboard Overview', level=1)
    doc.add_paragraph("The dashboard provides real-time counts for events and keeps you informed about upcoming activities.")
    
    # --- 4. Site Features ---
    doc.add_heading('4. Site Features & Scraper Settings', level=1)
    doc.add_paragraph("Manage the homepage carousel and scraper configuration.")
    if os.path.exists(SCREENSHOT_PATHS["features"]):
        doc.add_picture(SCREENSHOT_PATHS["features"], width=Inches(6))
        cap = doc.add_paragraph("Figure 2: Site Features and Robot Scraper settings.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 5. Policies ---
    doc.add_heading('5. Policies Management', level=1)
    doc.add_paragraph("The policies module allows archiving Circulars, Memos, and Local Government Policies with dedicated years and categories.")
    if os.path.exists(SCREENSHOT_PATHS["policies"]):
        doc.add_picture(SCREENSHOT_PATHS["policies"], width=Inches(6))
        cap = doc.add_paragraph("Figure 3: GAD Policy Management portal.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 6. Events ---
    doc.add_heading('6. Calendar Management', level=1)
    doc.add_paragraph("Maintain a current list of activities for the public calendar.")
    if os.path.exists(SCREENSHOT_PATHS["events"]):
        doc.add_picture(SCREENSHOT_PATHS["events"], width=Inches(6))
        cap = doc.add_paragraph("Figure 4: Calendar Event Creation Tool.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 7. Projects ---
    doc.add_heading('7. Project Archives', level=1)
    doc.add_paragraph("Track project progress across Proposed, Ongoing, and Completed stages.")
    if os.path.exists(SCREENSHOT_PATHS["projects"]):
        doc.add_picture(SCREENSHOT_PATHS["projects"], width=Inches(6))
        cap = doc.add_paragraph("Figure 5: GAD Project Listing dashboard.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 8. Knowledge Products ---
    doc.add_heading('8. Knowledge Products', level=1)
    doc.add_paragraph("Distribute official reports and learning materials.")
    if os.path.exists(SCREENSHOT_PATHS["knowledge"]):
        doc.add_picture(SCREENSHOT_PATHS["knowledge"], width=Inches(6))
        cap = doc.add_paragraph("Figure 6: Learning materials management.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 9. Brochures ---
    doc.add_heading('9. Brochures', level=1)
    if os.path.exists(SCREENSHOT_PATHS["brochures"]):
        doc.add_picture(SCREENSHOT_PATHS["brochures"], width=Inches(6))
        cap = doc.add_paragraph("Figure 7: Manage Digital Brochures.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 10. Livelihood ---
    doc.add_heading('10. Livelihood Program feeds', level=1)
    if os.path.exists(SCREENSHOT_PATHS["livelihood"]):
        doc.add_picture(SCREENSHOT_PATHS["livelihood"], width=Inches(6))
        cap = doc.add_paragraph("Figure 8: Integrating social media feeds.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 11. Org Structure ---
    doc.add_heading('11. Org Structure & GFPS Committee', level=1)
    doc.add_paragraph("Update the institutional hierarchy and committee member profiles.")
    if os.path.exists(SCREENSHOT_PATHS["org"]):
        doc.add_picture(SCREENSHOT_PATHS["org"], width=Inches(6))
        cap = doc.add_paragraph("Figure 9: Hierarchy management.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(SCREENSHOT_PATHS["committee"]):
        doc.add_picture(SCREENSHOT_PATHS["committee"], width=Inches(6))
        cap = doc.add_paragraph("Figure 10: Committee member updates.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 12. Tracking ---
    doc.add_heading('12. Tracking Matrix', level=1)
    doc.add_paragraph("Real-time logging of all administrative actions for audit compliance.")
    if os.path.exists(SCREENSHOT_PATHS["tracking"]):
        doc.add_picture(SCREENSHOT_PATHS["tracking"], width=Inches(6))
        cap = doc.add_paragraph("Figure 11: System Audit Log and Tracking Matrix.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\nEnd of Document")

    # --- Save ---
    output_path = os.path.join(os.getcwd(), "GAD_Admin_User_Manual.docx")
    doc.save(output_path)
    print(f"Update: DOCX generated successfully at: {output_path}")

if __name__ == "__main__":
    create_manual()
